import os
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


class WeeklyReportGenerator:
    def __init__(self, output_dir: str = "reports"):
        self.output_dir = output_dir
        self.ensure_output_dir()
        self.setup_fonts()

    def ensure_output_dir(self):
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def setup_fonts(self):
        try:
            font_paths = [
                "/Users/hwangeunbi/Documents/GitHub/report-generator/font/PretendardVariable.ttf",
            ]
            font_registered = False
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont("Korean", font_path))
                        font_registered = True
                        break
                    except Exception as e:
                        continue
            if not font_registered:
                self.download_nanum_font()
        except Exception:
            pass

    def download_nanum_font(self):
        try:
            import urllib.request

            font_dir = os.path.join(self.output_dir, "fonts")
            if not os.path.exists(font_dir):
                os.makedirs(font_dir)
            font_path = os.path.join(font_dir, "NanumGothic.ttf")
            if not os.path.exists(font_path):
                font_url = "https://github.com/naver/nanumfont/raw/master/fonts/NanumGothic.ttf"
                urllib.request.urlretrieve(font_url, font_path)
            pdfmetrics.registerFont(TTFont("Korean", font_path))
        except Exception:
            pass

    def generate_report_data(
        self,
        week_number: int,
        week_input: int,
        summaries: Dict[str, str],
        start_date: str,
        end_date: str,
        student: str = "황은비",
        author: str = "하마랩",
    ) -> Dict[str, Any]:
        return {
            "title": f"현장실습 주간보고서 ( {week_input} 주차)",
            "period": "2025.07.01 - 2025.12.21",
            "company": "하마랩",
            "student": student,
            "week_number": week_number,
            "week_input": week_input,
            "sections": [
                {
                    "number": "1",
                    "title": "주간 진행 사항",
                    "content": summaries.get("weekly", "내용을 입력하세요."),
                },
                {
                    "number": "2",
                    "title": "문제 및 대책",
                    "content": summaries.get("problem", "내용을 입력하세요."),
                },
                {
                    "number": "3",
                    "title": "소감",
                    "content": summaries.get("thoughts", "내용을 입력하세요."),
                },
                {
                    "number": "4",
                    "title": "차주 계획",
                    "content": summaries.get("plan", "내용을 입력하세요."),
                },
            ],
            "signature_line": "기술지도위원 확인",
            "signature_space": "(인)",
        }

    def create_word_report(self, report_data: Dict[str, Any]) -> str:
        doc = Document()
        doc.core_properties.language = "ko-KR"
        title_para = doc.add_heading(level=1)
        title_run = title_para.add_run(report_data["title"])
        title_run.font.name = "맑은 고딕"
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cells = table.rows[0].cells
        cells[0].text = report_data["period"]
        cells[0].merge(cells[1])
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cells = table.rows[1].cells
        cells[0].text = f"근무 회사 및 부서명: {report_data['company']}"
        cells[1].text = f"학생명: {report_data['student']}"

        cells = table.rows[2].cells
        cells[0].text = "현장실습 활동사항"
        cells[0].merge(cells[1])
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cells = table.rows[3].cells
        cells[0].text = ""
        cells[0].merge(cells[1])

        for section in report_data["sections"]:
            para = cells[0].add_paragraph()
            heading_run = para.add_run(f"{section['number']}. {section['title']}\n")
            heading_run.bold = True
            heading_run.font.name = "맑은 고딕"
            heading_run.font.size = Pt(11)
            content_run = para.add_run(section["content"])
            content_run.font.name = "맑은 고딕"
            content_run.font.size = Pt(10)

        doc.add_paragraph()
        signature_table = doc.add_table(rows=2, cols=1)
        signature_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        signature_table.rows[0].cells[0].text = report_data["signature_line"]
        signature_table.rows[1].cells[0].text = report_data["signature_space"]

        filename = f"주간보고서_{report_data['week_input']}주차_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def create_pdf_report(self, report_data: Dict[str, Any]) -> str:
        """PDF 보고서 생성"""
        filename = f"주간보고서_{report_data['week_input']}주차_{datetime.now().strftime('%Y%m%d')}.pdf"
        filepath = os.path.join(self.output_dir, filename)

        doc = SimpleDocTemplate(filepath, pagesize=A4)
        story = []

        styles = getSampleStyleSheet()

        # 스타일 정의
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontName=(
                "Korean"
                if "Korean" in pdfmetrics.getRegisteredFontNames()
                else "Helvetica"
            ),
            fontSize=16,
            alignment=1,
            spaceAfter=20,
        )

        normal_style = ParagraphStyle(
            "CustomNormal",
            parent=styles["Normal"],
            fontName=(
                "Korean"
                if "Korean" in pdfmetrics.getRegisteredFontNames()
                else "Helvetica"
            ),
            fontSize=10,
            spaceAfter=6,
            leading=13,
            leftIndent=0,
        )

        # 제목
        story.append(Paragraph(report_data["title"], title_style))
        story.append(Spacer(1, 20))

        # 섹션 내용을 하나의 HTML 문자열로 구성
        section_html = ""
        for section in report_data["sections"]:
            section_html += f"<b>{section['number']}. {section['title']}</b><br/>{section['content'].replace(chr(10), '<br/>')}<br/><br/>"

        section_paragraph = Paragraph(section_html.strip(), normal_style)

        # 표 데이터 구성
        table_data = [
            [report_data["period"], ""],
            [
                f"근무 회사 및 부서명: {report_data['company']}",
                f"학생명: {report_data['student']}",
            ],
            ["현장실습 활동사항", ""],
            [section_paragraph, ""],  # ✅ 섹션을 Paragraph로 삽입
        ]

        table = Table(
            table_data,
            colWidths=[3 * inch, 3 * inch],
            rowHeights=[
                0.5 * inch,
                0.5 * inch,
                0.5 * inch,
                None,
            ],  # 마지막 row height은 자동
        )
        table.setStyle(
            TableStyle(
                [
                    (
                        "FONTNAME",
                        (0, 0),
                        (-1, -1),
                        (
                            "Korean"
                            if "Korean" in pdfmetrics.getRegisteredFontNames()
                            else "Helvetica"
                        ),
                    ),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("SPAN", (0, 0), (1, 0)),  # 기간
                    ("SPAN", (0, 2), (1, 2)),  # 활동사항 타이틀
                    ("SPAN", (0, 3), (1, 3)),  # 섹션 내용
                ]
            )
        )

        story.append(table)
        story.append(Spacer(1, 30))

        # 서명란
        signature_table = Table(
            [[report_data["signature_line"]], [report_data["signature_space"]]],
            colWidths=[2 * inch],
        )
        signature_table.setStyle(
            TableStyle(
                [
                    (
                        "FONTNAME",
                        (0, 0),
                        (-1, -1),
                        (
                            "Korean"
                            if "Korean" in pdfmetrics.getRegisteredFontNames()
                            else "Helvetica"
                        ),
                    ),
                    ("ALIGN", (0, 0), (-1, -1), "RIGHT"),
                ]
            )
        )
        story.append(signature_table)

        doc.build(story)
        return filepath

    def generate_reports(
        self,
        week_number: int,
        week_input: int,
        summaries: Dict[str, str],
        start_date: str,
        end_date: str,
        student: str = "황은비",
        author: str = "하마랩",
    ) -> Dict[str, str]:
        report_data = self.generate_report_data(
            week_number, week_input, summaries, start_date, end_date, student, author
        )
        try:
            word_path = self.create_word_report(report_data)
            print(f"✅ Word 보고서 생성 완료: {word_path}")
        except Exception as e:
            print(f"❌ Word 보고서 생성 실패: {e}")
            word_path = None

        try:
            pdf_path = self.create_pdf_report(report_data)
            print(f"✅ PDF 보고서 생성 완료: {pdf_path}")
        except Exception as e:
            print(f"❌ PDF 보고서 생성 실패: {e}")
            pdf_path = None

        return {"word": word_path, "pdf": pdf_path}
